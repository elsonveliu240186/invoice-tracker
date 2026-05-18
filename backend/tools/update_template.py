"""
One-off script: replace literal placeholder text in invoice-template.docx
with poi-tl token equivalents so the PoiTlInvoiceDocxRenderer can fill them.

Run from the backend/ directory:
    python tools/update_template.py

Table 0 (company + invoice header):
  Row 0 cell 0: [Company Name]    -> {{companyName}}
  Row 1 cell 0: [Street Address]  -> {{companyAddress}}
  Row 2 cell 0: [City, ST  ZIP]   -> {{companyPhone}}
  Row 3 cell 1: INVOICE #         -> INVOICE #   (keep label)
  Row 3 cell 2: DATE              -> DATE        (keep label)
  Row 4 cell 1: 2034 (number)     -> {{invoiceNumber}}
  Row 4 cell 2: 5/1/2014 (date)   -> {{invoiceIssueDate}}
  Row 6 cell 1: 564               -> {{invoiceDueDate}} (use due date for TERMS row val)

Table 1 (bill to / ship to):
  Row 1 cell 0: [Name]            -> {{clientName}}
  Row 2 cell 0: [Company Name]    -> {{clientAddress}}
  Row 1 cell 2: [Name]            -> {{clientName}}   (ship to - same client)
  Row 2 cell 2: [Company Name]    -> (blank - simplify)

Table 2 (line items) — restructure for {{lines}} trigger:
  Row 0: keep header (DESCRIPTION, QTY, UNIT PRICE, AMOUNT)
  Row 1: becomes trigger row: {{lines}}, empty, empty, empty
  Row 2: becomes template row: {{description}}, {{quantity}}, {{unitPrice}}, {{lineTotal}}
  Remaining rows: remove (replaced at runtime by PoiTlInvoiceDocxRenderer)
  Keep last row (TOTAL row) with {{invoiceSubtotal}}/{{invoiceTotal}} etc.

After tables, add footer paragraph with email/VAT info using tokens.
"""
import shutil
import re
from docx import Document
from docx.oxml.ns import qn

TEMPLATE_PATH = (
    r"C:\Users\ExpertBook\agenticai\projects\invoice-tracker"
    r"\backend\src\main\resources\templates\invoice-template.docx"
)
BACKUP_PATH = TEMPLATE_PATH + ".bak"


def set_cell_text(cell, text):
    """Replace all text runs in the first paragraph of a cell with a single run."""
    para = cell.paragraphs[0]
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)


def remove_table_row(table, row_index):
    """Remove a row from the table by its index."""
    tr = table.rows[row_index]._tr
    table._tbl.remove(tr)


def main():
    shutil.copy2(TEMPLATE_PATH, BACKUP_PATH)
    print(f"Backup saved to: {BACKUP_PATH}")

    doc = Document(TEMPLATE_PATH)

    # ---- Table 0: Company header + invoice metadata ----
    t0 = doc.tables[0]
    set_cell_text(t0.rows[0].cells[0], "{{companyName}}")
    set_cell_text(t0.rows[1].cells[0], "{{companyAddress}}")
    set_cell_text(t0.rows[2].cells[0], "{{companyPhone}}")
    # Row 4: invoice number and issue date
    set_cell_text(t0.rows[4].cells[1], "{{invoiceNumber}}")
    set_cell_text(t0.rows[4].cells[2], "{{invoiceIssueDate}}")
    # Row 6 cell 2: due date value
    set_cell_text(t0.rows[6].cells[2], "{{invoiceDueDate}}")

    # ---- Table 1: Bill-to / Ship-to ----
    t1 = doc.tables[1]
    set_cell_text(t1.rows[1].cells[0], "{{clientName}}")
    set_cell_text(t1.rows[2].cells[0], "{{clientAddress}}")
    set_cell_text(t1.rows[3].cells[0], "{{clientEmail}}")
    # Clear the ship-to side (we only track one address)
    set_cell_text(t1.rows[1].cells[2], "")
    set_cell_text(t1.rows[2].cells[2], "")
    set_cell_text(t1.rows[3].cells[2], "")
    set_cell_text(t1.rows[4].cells[0], "")
    set_cell_text(t1.rows[4].cells[2], "")
    set_cell_text(t1.rows[5].cells[0], "")
    set_cell_text(t1.rows[5].cells[2], "")
    set_cell_text(t1.rows[6].cells[0], "")
    set_cell_text(t1.rows[6].cells[1], "")
    set_cell_text(t1.rows[6].cells[2], "")

    # ---- Table 2: Line items + totals ----
    t2 = doc.tables[2]
    # Row 0: header stays (DESCRIPTION, QTY, UNIT PRICE, AMOUNT)

    # Transform Row 1 -> trigger row: {{lines}} in first cell
    set_cell_text(t2.rows[1].cells[0], "{{lines}}")
    set_cell_text(t2.rows[1].cells[1], "")
    set_cell_text(t2.rows[1].cells[2], "")
    set_cell_text(t2.rows[1].cells[3], "")

    # Transform Row 2 -> template row: {{description}}, {{quantity}}, {{unitPrice}}, {{lineTotal}}
    set_cell_text(t2.rows[2].cells[0], "{{description}}")
    set_cell_text(t2.rows[2].cells[1], "{{quantity}}")
    set_cell_text(t2.rows[2].cells[2], "{{unitPrice}}")
    set_cell_text(t2.rows[2].cells[3], "{{lineTotal}}")

    # The last row (index 19) has the TOTAL - update it with tokens
    last_row_idx = len(t2.rows) - 1
    last_row = t2.rows[last_row_idx]
    set_cell_text(last_row.cells[0], "{{invoiceSubtotal}} (subtotal)  Tax ({{invoiceTaxRate}}): {{invoiceTaxAmount}}")
    set_cell_text(last_row.cells[1], "TOTAL")
    set_cell_text(last_row.cells[2], "")
    set_cell_text(last_row.cells[3], "{{invoiceTotal}}")

    # Remove rows 3 through last_row_idx-1 (the filler \xa0 rows)
    # Remove from the end to avoid index shifting
    rows_to_remove = list(range(3, last_row_idx))
    for idx in reversed(rows_to_remove):
        remove_table_row(t2, idx)

    # ---- Paragraph 4: update the contact line ----
    # Paragraph index 4 is '[Name, Phone, email@address.com]' — replace with tokens
    p4 = doc.paragraphs[4]
    for run in p4.runs:
        run.text = ""
    if p4.runs:
        p4.runs[0].text = "{{companyName}} | {{companyPhone}} | {{companyEmail}}"
    else:
        p4.add_run("{{companyName}} | {{companyPhone}} | {{companyEmail}}")

    doc.save(TEMPLATE_PATH)
    print(f"Template updated: {TEMPLATE_PATH}")

    # Verify
    doc2 = Document(TEMPLATE_PATH)
    print("\nVerification - Table 0:")
    for r_idx, row in enumerate(doc2.tables[0].rows):
        cells = [cell.text for cell in row.cells]
        print(f"  Row {r_idx}: {cells}")
    print("\nVerification - Table 2:")
    for r_idx, row in enumerate(doc2.tables[2].rows):
        cells = [cell.text for cell in row.cells]
        print(f"  Row {r_idx}: {cells}")


if __name__ == "__main__":
    main()
